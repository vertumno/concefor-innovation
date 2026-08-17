#!/usr/bin/env python3
"""Gera os sete roteiros editáveis do Concefor Podcast em DOCX.

Design: preset compact_reference_guide + cabeçalho workshop_agenda.
Overrides nomeados: página A4 e densidade spoken_script_two_page.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips


NAVY = "173A6B"
TEAL = "008EA3"
GOLD = "D89900"
INK = "20242B"
MUTED = "5E6875"
PALE_TEAL = "EAF7F8"
PALE_GOLD = "FFF7E3"
PALE_BLUE = "EEF3F9"
LIGHT_GRAY = "D9E0E8"
VERY_LIGHT_GRAY = "F7F9FB"
WHITE = "FFFFFF"
FONT = "Calibri"


EPISODES = [
    {
        "ep": 1,
        "source": "ep1-vanessa-battestin.md",
        "output": "roteiro-ep1-vanessa-battestin.docx",
        "extras": [
            "Quando financiamento e avaliação aparecem como lacunas, qual deles tende a travar o outro: sem financiamento não se avalia bem ou, sem evidências, não se conquista financiamento?",
            "Que indicador de qualidade pode parecer convincente em um relatório, mas dizer pouco sobre aprendizagem e permanência?",
            "Numa Rede Federal tão diversa, o que precisa ser política comum e o que deve continuar sob decisão de cada campus?",
            "O que o Cefor pode adaptar das grandes universidades de educação a distância - e o que deveria recusar copiar?",
            "Nos MOOCs trilíngues, traduzir é apenas uma camada. Pedagogicamente, o que torna um curso culturalmente relevante?",
            "O que faz uma parceria internacional produzir mudança pedagógica real, em vez de apenas ampliar a circulação institucional?",
            "Escala, acessibilidade e presença territorial podem puxar o curso em direções diferentes. Que decisão de desenho ajuda a conciliá-las?",
            "Num curso aberto para muitas pessoas, que forma de presença da instituição não pode desaparecer em nome da escala?",
            "Que dados seria importante acompanhar numa próxima oferta para compreender as trajetórias dos participantes sem reduzi-las ao rótulo de evasão?",
            "Entre os aprendizados desses vinte anos, qual não deve virar apenas memória porque ainda precisa orientar o trabalho cotidiano?",
        ],
    },
    {
        "ep": 2,
        "source": "ep2-jaqueline-sanz.md",
        "output": "roteiro-ep2-jaqueline-sanz.docx",
        "extras": [
            "Quando os números e a escuta das mulheres apontam prioridades diferentes, como uma equipe pública deve decidir o que pesa mais?",
            "O Atlas reúne mulheres de 19 segmentos. Você pode dar um exemplo de solução que parece inclusiva quando pensamos numa usuária abstrata, mas falha diante de um território concreto?",
            "O que uma planilha não consegue mostrar e precisa ser preservado quando um relato vira diagnóstico ou política pública?",
            "Dados mais detalhados aumentam a visibilidade, mas também podem expor grupos pequenos. Como equilibrar precisão, proteção e participação?",
            "Quase metade das participantes do Atlas são as principais provedoras de suas famílias. O que muda no desenho de um serviço quando trabalho remunerado e cuidado precisam ser considerados juntos?",
            "Em 'Conexões Quilombolas', memória, território e tecnologias tradicionais aparecem lado a lado. O que muda quando a inovação começa reconhecendo o conhecimento que já existe na comunidade?",
            "Antes de substituir um atendimento presencial por um aplicativo, qual pergunta uma instituição deveria ser obrigada a responder?",
            "Ao transformar pesquisa em exposição e podcast, como evitar que histórias complexas virem apenas ilustrações emocionais?",
            "Além de acessos, inscrições ou downloads, que sinal mostraria que uma tecnologia pública realmente melhorou a vida das mulheres?",
            "Daqui a dois anos, que mudança concreta faria você dizer que os dados e as conversas do Atlas produziram consequência real?",
        ],
    },
    {
        "ep": 3,
        "source": "ep3-mariella-berger.md",
        "output": "roteiro-ep3-mariella-berger.docx",
        "extras": [
            "O que torna um caso suficientemente aberto para provocar julgamento profissional, mas claro o bastante para não desorientar quem estuda a distância?",
            "Um caso muito próximo da realidade também pode sobrecarregar o estudante com informações. Como calibrar complexidade e foco?",
            "Que intervenção do professor transforma divergência entre estudantes em raciocínio, sem antecipar a resposta?",
            "Como avaliar uma decisão construída pelo método de casos quando existem várias respostas defensáveis?",
            "Na era das plataformas que geram conteúdo e devolutivas, qual parte do trabalho docente se torna ainda mais indispensável?",
            "Que tipo de tecnologia deveria ser retirada de um curso quando aumenta a interação, mas não melhora a aprendizagem?",
            "No novo marco regulatório da EaD, que exigência protege efetivamente a qualidade e qual corre o risco de produzir apenas burocracia?",
            "Que evidência um processo regulatório deveria pedir para verificar qualidade sem impor um único modelo pedagógico?",
            "Da experiência de gestão durante a pandemia, qual solução emergencial mereceu permanecer e qual não deveria ter sido normalizada?",
            "Antes de acrescentar um novo recurso à sala Moodle, que pergunta o professor deveria fazer para saber se ele realmente contribui para a aprendizagem?",
        ],
    },
    {
        "ep": 4,
        "source": "ep4-felipe-tessarolo.md",
        "output": "roteiro-ep4-felipe-tessarolo.docx",
        "extras": [
            "Seu trabalho sobre a primeira atividade avaliada mostra que esse começo pode ser decisivo na EaD. Que apoio deveria chegar antes que a dificuldade vire abandono?",
            "Se um projeto de IA aumenta engajamento, mas ainda não demonstra aprendizagem, o que a instituição deve fazer: ampliar, manter como piloto ou interromper?",
            "Ao testar uma tecnologia educacional, qual grupo de estudantes costuma revelar primeiro as barreiras que o desenho deixou passar?",
            "No projeto ADMIT, vocês trabalham com formação docente para uso responsável de IA. Qual competência mínima todo professor deveria desenvolver antes de levar essa tecnologia para uma turma?",
            "Como incluir estudantes no desenho e na avaliação de uma IA sem transformar essa participação em uma consulta simbólica?",
            "Quando uma resposta parece correta e bem escrita, como o estudante pode perceber que perspectivas, referências ou vozes ficaram ausentes?",
            "A padronização costuma ser tratada como inimiga da personalização. Em que situações ela ainda protege qualidade e equidade?",
            "O que uma instituição brasileira de EaD pode aprender com a Open University sem copiar soluções que dependem de outra infraestrutura e outro perfil de estudante?",
            "Em uma educação cada vez mais automatizada, que forma de presença humana não pode ser reduzida a atendimento ou suporte técnico?",
            "Se voltarmos a esta conversa em cinco anos, qual evidência mostraria que a IA realmente transformou a aprendizagem - e não apenas acelerou processos?",
        ],
    },
    {
        "ep": 5,
        "source": "ep5-rutinelli-favero.md",
        "output": "roteiro-ep5-rutinelli-favero.docx",
        "extras": [
            "O que uma comunidade de prática contínua, como o Papo com IA.IÁ, consegue revelar que uma formação isolada dificilmente mostraria?",
            "Como transformar dúvidas e descobertas compartilhadas nesses encontros em decisões institucionais, sem perder a abertura da conversa?",
            "O que muda quando os estudantes participam da construção dos critérios éticos para uso de IA, em vez de apenas receberem uma regra pronta?",
            "Como fazer um manual de uso ético continuar útil quando as ferramentas mudam tão rapidamente? Que princípios resistem a essa mudança?",
            "Em que etapa do design educacional o ChatGPT pode ampliar alternativas e em que etapa tende a homogeneizar a proposta?",
            "Você poderia dar um exemplo de resultado produzido com IA que parece eficiente, mas representa uma decisão pedagógica ruim?",
            "Antes de publicar uma atividade planejada com apoio de IA, que revisão humana deveria ser obrigatória?",
            "Para quem está chegando agora a um NTE, o que deveria vir primeiro: domínio da ferramenta, julgamento pedagógico ou acessibilidade?",
            "Qual decisão educacional nunca deveria ser automatizada, mesmo que futuramente a tecnologia consiga executá-la?",
            "Que hábito de formação docente precisa virar rotina antes da chegada da próxima ferramenta?",
        ],
    },
    {
        "ep": 6,
        "source": "ep6-mariano-pimentel.md",
        "output": "roteiro-ep6-mariano-pimentel.docx",
        "extras": [
            "O livro fala em 'leitura generativa'. O que o estudante precisa fazer para que a IA amplie o encontro com o texto, em vez de apenas encurtá-lo?",
            "No experimento do GPT Paulo Freire, o que uma simulação pode ajudar a pensar e o que ela inevitavelmente distorce?",
            "Como distinguir uma colaboração intelectual de uma conversa apenas convincente com um chatbot?",
            "Que atividade de aprendizagem fica pior - e não melhor - quando a IA acelera sua execução?",
            "O livro discute o risco de substituição da docência. Que julgamento pedagógico você considera inegociavelmente humano?",
            "Entre proibir e liberar qualquer uso, que terceira via uma instituição pode adotar para orientar estudantes e professores?",
            "Se saber conversar com a IA virar uma vantagem acadêmica, como evitar que essa habilidade produza uma nova desigualdade entre estudantes?",
            "Quando a IA responde com segurança, mas está errada, como transformar a verificação em aprendizagem, e não apenas em correção?",
            "Antes de escalar um assistente de IA para milhares de estudantes, que evidência um projeto-piloto deveria produzir?",
            "Que prática hoje apresentada como inovação você interromperia até que as instituições façam perguntas melhores?",
        ],
    },
    {
        "ep": 7,
        "source": "ep7-mauro-oliveira.md",
        "output": "roteiro-ep7-mauro-oliveira.docx",
        "extras": [
            "Como redesenhar uma avaliação para que o estudante possa usar IA e, ainda assim, precise demonstrar compreensão, julgamento e autoria?",
            "Em que momento a ferramenta deixa de apoiar o pensamento e passa a substituí-lo? Que sinal concreto ajuda a perceber essa mudança?",
            "Que atividade simples um professor pode aplicar para transformar uma resposta pronta da IA em ponto de partida para investigação e debate?",
            "Quando uma IA responde com muita segurança e pouca base, que hábito intelectual a escola precisa ensinar para que o estudante não confunda fluência com verdade?",
            "Soberania digital também passa pelas condições materiais dos datacenters. Como energia, água e capacidade computacional deveriam entrar no letramento sobre IA?",
            "Nas conversas da Caravana de Soberania Digital, qual preocupação aparece com mais força dentro das universidades e ainda recebe pouca atenção pública?",
            "Como medir o impacto de um projeto de inclusão tecnológica sem se limitar a matrículas, certificados ou número de pessoas atendidas?",
            "Há quem argumente que desenvolver infraestrutura e soluções próprias custa caro demais. Em que situação usar uma plataforma estrangeira é uma escolha legítima, e quando vira dependência?",
            "Escolas com menos infraestrutura podem ficar ainda mais distantes das instituições que adotam IA rapidamente. Qual política pública evitaria que a inovação ampliasse essa desigualdade?",
            "Se uma instituição pública pudesse executar apenas um experimento educacional com IA nos próximos noventa dias, qual você escolheria e que resultado exigiria antes de continuar?",
        ],
    },
]


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def clean_md(value: str) -> str:
    value = re.sub(r"\*\*(.*?)\*\*", r"\1", value)
    value = re.sub(r"\*(.*?)\*", r"\1", value)
    value = value.replace("—", "-").replace("–", "-").replace("‑", "-")
    return re.sub(r"\s+", " ", value).strip()


def parse_script(path: Path) -> dict:
    raw = path.read_text(encoding="utf-8")
    first_line = clean_md(raw.splitlines()[0])
    title_match = re.match(r"^#\s+Roteiro\s+·\s+Ep\.\s+(\d+)\s+-\s+(.+)$", first_line)
    if not title_match:
        raise ValueError(f"Título não reconhecido em {path}")

    def metadata(label: str) -> str:
        match = re.search(rf"\|\s*\*\*{re.escape(label)}\*\*\s*\|\s*(.*?)\s*\|", raw)
        if not match:
            raise ValueError(f"Metadado {label} ausente em {path}")
        return clean_md(match.group(1))

    opening_section = raw.split("## Abertura e introdução do host", 1)[1].split("## Conversa", 1)[0]
    opening_blocks = [clean_md(block) for block in re.split(r"\n\s*\n", opening_section.strip()) if block.strip()]
    opening = re.sub(r"^Frase de abertura:\s*", "", opening_blocks[0], flags=re.I)
    introduction = opening_blocks[1]

    conversation = raw.split("## Conversa", 1)[1].split("## Considerações finais do host", 1)[0]
    blocks = [block.strip() for block in re.split(r"\n\s*\n", conversation.strip()) if block.strip()]
    questions: list[str] = []
    hooks: dict[int, str] = {}
    for block in blocks:
        normalized = clean_md(block)
        question_match = re.match(r"^(\d+)\.\s+(.*)$", normalized)
        if question_match:
            questions.append(question_match.group(2).strip())
            continue
        hook_match = re.match(r"^Gancho:\s*(.*)$", normalized, flags=re.I)
        if hook_match and questions:
            hooks[len(questions)] = hook_match.group(1).strip()

    if len(questions) != 5:
        raise ValueError(f"Esperadas 5 perguntas em {path}; encontradas {len(questions)}")

    closing = clean_md(raw.split("## Considerações finais do host", 1)[1])
    return {
        "episode": int(title_match.group(1)),
        "guest": title_match.group(2).strip(),
        "recording": metadata("Gravação"),
        "host": metadata("Host"),
        "event_talk": metadata("Fala no evento"),
        "opening": opening,
        "introduction": introduction,
        "questions": questions,
        "hooks": hooks,
        "closing": closing,
    }


def set_run_font(run, *, size=None, color=None, bold=None, italic=None, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def ensure_language(style, language="pt-BR"):
    rpr = style.element.get_or_add_rPr()
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), language)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05
    ensure_language(normal)

    title = styles["Title"]
    title.font.name = FONT
    title._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    title._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    title.font.size = Pt(21)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    ensure_language(title)

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT
    subtitle._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    subtitle._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    subtitle.font.size = Pt(9.5)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(6)
    subtitle.paragraph_format.keep_with_next = True
    ensure_language(subtitle)

    h1 = styles["Heading 1"]
    h1.font.name = FONT
    h1._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    h1._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    h1.font.size = Pt(17)
    h1.font.bold = True
    h1.font.color.rgb = rgb(NAVY)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(5)
    h1.paragraph_format.keep_with_next = True
    ensure_language(h1)

    h2 = styles["Heading 2"]
    h2.font.name = FONT
    h2._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    h2._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    h2.font.size = Pt(10)
    h2.font.bold = True
    h2.font.color.rgb = rgb(TEAL)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(2)
    h2.paragraph_format.keep_with_next = True
    ensure_language(h2)

    if "Kicker" not in styles:
        styles.add_style("Kicker", WD_STYLE_TYPE.PARAGRAPH)
    kicker = styles["Kicker"]
    kicker.font.name = FONT
    kicker._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    kicker._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    kicker.font.size = Pt(8.5)
    kicker.font.bold = True
    kicker.font.color.rgb = rgb(TEAL)
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(1)
    kicker.paragraph_format.keep_with_next = True
    ensure_language(kicker)

    if "Question" not in styles:
        styles.add_style("Question", WD_STYLE_TYPE.PARAGRAPH)
    question = styles["Question"]
    question.base_style = normal
    question.font.name = FONT
    question._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    question._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    question.font.size = Pt(10.1)
    question.font.color.rgb = rgb(INK)
    question.paragraph_format.space_before = Pt(0)
    question.paragraph_format.space_after = Pt(4)
    question.paragraph_format.line_spacing = 1.03
    question.paragraph_format.keep_together = True
    ensure_language(question)

    if "Extra Question" not in styles:
        styles.add_style("Extra Question", WD_STYLE_TYPE.PARAGRAPH)
    extra = styles["Extra Question"]
    extra.base_style = normal
    extra.font.name = FONT
    extra._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    extra._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    extra.font.size = Pt(10.2)
    extra.font.color.rgb = rgb(INK)
    extra.paragraph_format.space_before = Pt(0)
    extra.paragraph_format.space_after = Pt(5)
    extra.paragraph_format.line_spacing = 1.06
    extra.paragraph_format.keep_together = True
    ensure_language(extra)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=LIGHT_GRAY, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def add_metadata_cell(cell, label: str, value: str):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    label_run = p.add_run(label.upper())
    set_run_font(label_run, size=7.2, color=TEAL, bold=True)
    value_run = p.add_run("\n" + value)
    set_run_font(value_run, size=8.8, color=INK, bold=True)


def add_metadata_strip(doc, data, apply_table_geometry, column_widths_from_weights, content_width):
    table = doc.add_table(rows=1, cols=3)
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    row_properties.append(table_header)
    widths = column_widths_from_weights([1.35, 1.0, 3.35], content_width)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=content_width,
        indent_dxa=120,
        cell_margins_dxa={"top": 85, "bottom": 85, "start": 120, "end": 120},
    )
    set_table_borders(table)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, VERY_LIGHT_GRAY)
    add_metadata_cell(table.cell(0, 0), "Gravação", data["recording"])
    add_metadata_cell(table.cell(0, 1), "Host", data["host"])
    add_metadata_cell(table.cell(0, 2), "Fala no evento", data["event_talk"])
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(1)
    after.add_run("")


def set_paragraph_callout(paragraph, *, fill: str, border: str, left=120, right=100):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)

    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left_border = borders.find(qn("w:left"))
    if left_border is None:
        left_border = OxmlElement("w:left")
        borders.append(left_border)
    left_border.set(qn("w:val"), "single")
    left_border.set(qn("w:sz"), "18")
    left_border.set(qn("w:space"), "6")
    left_border.set(qn("w:color"), border)
    paragraph.paragraph_format.left_indent = Twips(left)
    paragraph.paragraph_format.right_indent = Twips(right)
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_together = True


def add_callout(doc, label: str, text: str, *, fill: str, border: str, italic=False, compact=False):
    p = doc.add_paragraph()
    set_paragraph_callout(p, fill=fill, border=border)
    p.paragraph_format.line_spacing = 1.0 if compact else 1.05
    label_run = p.add_run(label.upper() + "\n")
    set_run_font(label_run, size=7.6, color=border, bold=True)
    body_run = p.add_run(text)
    set_run_font(body_run, size=8.8 if compact else 9.5, color=INK, italic=italic)
    return p


def add_field(paragraph, instruction: str, placeholder="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8, color=MUTED)


def configure_header_footer(section, data, content_width):
    header = section.header
    hp = header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.line_spacing = 1.0
    hp.paragraph_format.tab_stops.add_tab_stop(Twips(content_width), WD_TAB_ALIGNMENT.RIGHT)
    left = hp.add_run("CONCEFOR PODCAST")
    set_run_font(left, size=7.5, color=TEAL, bold=True)
    right = hp.add_run(f"\tEP. {data['episode']} · {data['guest']}")
    set_run_font(right, size=7.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.clear()
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.line_spacing = 1.0
    fp.paragraph_format.tab_stops.add_tab_stop(Twips(content_width), WD_TAB_ALIGNMENT.RIGHT)
    left = fp.add_run("Roteiro editável · priorizar a transcrição quando disponível")
    set_run_font(left, size=7.5, color=MUTED)
    right = fp.add_run("\tPágina ")
    set_run_font(right, size=7.5, color=MUTED)
    add_field(fp, "PAGE")
    middle = fp.add_run(" de ")
    set_run_font(middle, size=7.5, color=MUTED)
    add_field(fp, "NUMPAGES", placeholder="2")


def add_numbering_definition(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    # Reserve enough room for two-digit items so "10." never touches the text.
    tab.set(qn("w:pos"), "650")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "650")
    ind.set(qn("w:hanging"), "380")
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), NAVY)
    r_pr.append(color)
    bold = OxmlElement("w:b")
    r_pr.append(bold)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_question(doc, text: str, num_id: int, *, extra=False, keep_with_next=False):
    p = doc.add_paragraph(style="Extra Question" if extra else "Question")
    apply_numbering(p, num_id)
    p.paragraph_format.keep_with_next = keep_with_next
    marker_match = re.match(r"^\[([^\]]+)\]\s*(.*)$", text)
    if marker_match:
        marker = marker_match.group(1).replace("ATUALIZAR COM A PALESTRA", "PRIORIZAR TRANSCRIÇÃO")
        marker_run = p.add_run(marker + " - ")
        set_run_font(marker_run, size=8.8, color=GOLD, bold=True)
        text = marker_match.group(2)
    body = p.add_run(text)
    set_run_font(body, size=10.2 if extra else 10.1, color=INK)
    return p


def transcript_status(data) -> str:
    if "antes" in data["recording"].lower():
        return (
            "Situação deste episódio: a gravação acontece antes da fala no evento. Ajuste a primeira "
            "pergunta com o briefing do convidado; se houver captação posterior, considere uma retomada "
            "curta da palestra."
        )
    return (
        "Situação deste episódio: ainda não há transcrição no repositório. A pergunta marcada deve "
        "ser substituída por uma retomada concreta da palestra ou mesa assim que o texto estiver disponível."
    )


def build_docx(data, extras, output_path, table_helpers):
    apply_table_geometry, column_widths_from_weights, section_content_width_dxa, audit_docx_tables = table_helpers
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(6)
    content_width = section_content_width_dxa(section)
    configure_header_footer(section, data, content_width)

    doc.core_properties.title = f"Concefor Podcast - Ep. {data['episode']} - {data['guest']}"
    doc.core_properties.subject = "Roteiro editável do host"
    doc.core_properties.author = "Concefor Podcast"
    doc.core_properties.keywords = "Concefor, podcast, roteiro, entrevista"

    kicker = doc.add_paragraph("CONCEFOR PODCAST · ROTEIRO DO HOST", style="Kicker")
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title = doc.add_paragraph(f"EP. {data['episode']} · {data['guest']}", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = doc.add_paragraph("Versão editável para condução e impressão", style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_metadata_strip(doc, data, apply_table_geometry, column_widths_from_weights, content_width)

    transcript_rule = (
        "Se houver transcrição, reescreva de 4 a 6 perguntas principais a partir de afirmações, "
        "exemplos e tensões da fala. O roteiro deve partir principalmente da palestra, não apenas "
        "do currículo do convidado. Mova as perguntas substituídas para 'Carta na manga'; esse banco "
        "não tem limite. Em vez de pedir repetição, peça explicação, consequência, exemplo, contraponto ou decisão. "
        + transcript_status(data)
    )
    add_callout(
        doc,
        "Regra editorial · transcrição primeiro",
        transcript_rule,
        fill=PALE_GOLD,
        border=GOLD,
        compact=True,
    )

    doc.add_paragraph("ABERTURA", style="Heading 2")
    add_callout(doc, "Fala do host", data["opening"], fill=PALE_GOLD, border=GOLD)

    doc.add_paragraph("INTRODUÇÃO DO CONVIDADO", style="Heading 2")
    add_callout(doc, "Fala do host", data["introduction"], fill=PALE_TEAL, border=TEAL)

    doc.add_paragraph("CONVERSA PRINCIPAL · VERSÃO ATUAL SEM TRANSCRIÇÃO", style="Heading 2")
    main_num_id = add_numbering_definition(doc)
    for index, question in enumerate(data["questions"], start=1):
        has_hook = index in data["hooks"]
        add_question(doc, question, main_num_id, keep_with_next=has_hook)
        if has_hook:
            add_callout(
                doc,
                "Gancho · fala do host",
                data["hooks"][index],
                fill=PALE_GOLD,
                border=GOLD,
                italic=True,
                compact=True,
            )

    doc.add_paragraph("ENCERRAMENTO", style="Heading 2")
    closing = add_callout(doc, "Fala do host", data["closing"], fill=PALE_BLUE, border=NAVY)
    closing.paragraph_format.space_after = Pt(0)

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    page_break.paragraph_format.space_before = Pt(0)
    page_break.paragraph_format.space_after = Pt(0)

    doc.add_paragraph("CARTA NA MANGA", style="Kicker")
    doc.add_paragraph(f"Perguntas extras · {data['guest']}", style="Heading 1")
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(0)
    intro.paragraph_format.space_after = Pt(7)
    intro.paragraph_format.line_spacing = 1.05
    lead = intro.add_run("Como usar: ")
    set_run_font(lead, size=9.3, color=TEAL, bold=True)
    body = intro.add_run(
        "escolha uma ou duas conforme o rumo da conversa. Se a transcrição substituir perguntas "
        "da página 1, mova-as para esta lista. O banco pode crescer sem limite."
    )
    set_run_font(body, size=9.3, color=MUTED)

    extras_num_id = add_numbering_definition(doc)
    for question in extras:
        add_question(doc, question, extras_num_id, extra=True)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run("NOTAS DO HOST:  ________________________________________________________________")
    set_run_font(note_run, size=8.5, color=MUTED, bold=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    issues = audit_docx_tables(output_path)
    if issues:
        raise RuntimeError(f"Falha de geometria em {output_path.name}: {issues} problema(s)")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--skill-dir", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()

    scripts_dir = Path(__file__).resolve().parent.parent / "roteiros"
    sys.path.insert(0, str(args.skill_dir / "scripts"))
    from table_geometry import (  # pylint: disable=import-outside-toplevel
        apply_table_geometry,
        audit_docx_tables,
        column_widths_from_weights,
        section_content_width_dxa,
    )

    helpers = (
        apply_table_geometry,
        column_widths_from_weights,
        section_content_width_dxa,
        audit_docx_tables,
    )
    args.output_dir.mkdir(parents=True, exist_ok=True)
    for episode in EPISODES:
        data = parse_script(scripts_dir / episode["source"])
        if data["episode"] != episode["ep"]:
            raise ValueError(f"Episódio divergente em {episode['source']}")
        output_path = args.output_dir / episode["output"]
        build_docx(data, episode["extras"], output_path, helpers)
        print(output_path)


if __name__ == "__main__":
    main()
